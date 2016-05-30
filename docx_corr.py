from docx import Document
from string import *

d = Document('d:\\tmp\\K12\\01.02.16.docx')

i = 0
j = 0
k = 0

tabl = d.tables

lin = d.tables[1].rows[10].cells[0].text[-13:-3].split('.')
lin=lin[2]+lin[1]+lin[0]

print(lin)