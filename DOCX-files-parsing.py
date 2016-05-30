from HandlerConf import *
from docx import Document
import os


docxs = []
month = {
    'января': '01',
    'февраля': '02',
    'марта': '03',
    'апреля': '04',
    'мая': '05',
    'июня': '06',
    'июля': '07',
    'августа': '08',
    'сентября': '09',
    'октября': '10',
    'ноября': '11',
    'декабря': '12'
    }

sqlLine = al.text('drop table "Chim"; CREATE TABLE "Chim" ( "Time" TIMESTAMP, 	"EPropID" int,	"Value" FLOAT )')
try:
    q = con.execute(sqlLine)
    print('Chim создана!')
    sqlLine = ''
except:
    print('Не удалось создать Chim')
    exit()

outfile = open('d:\\tmp\\out.txt','w')

dir = 'd:\\tmp\\k12\\'


for name in os.listdir(dir):
    path = os.path.join(dir, name)
    if os.path.isfile(path):
        if path[-4:] != 'docx':
            pass
        else:
            docxs.append(path)

for dd in docxs:
    print('Parsed file:',dd)
    d = Document(dd)

    dateStr =  d.paragraphs[4].text[3:-7]
    dat = dateStr.split(' ')

    lin = d.tables[1].rows[10].cells[0].text[-13:-3].split('.')
    lin = lin[2] + lin[1] + lin[0]

    print(lin)

    dateStr = dat[3][0:4] + month.get(dat[1]) + dat[0][1:3]

    #outfile.write('\t'+dateStr)
    outfile.write('\t' + lin)
    tabs = d.tables
    epID = 0
    i=0
    value = 0

    for row in tabs[3].rows:
        j = 0
        for cell in row.cells:
            if i == 0:
                pass
            else:
                str = cell.text.lstrip()
                cr = str.find('\n')
                if cr > 0:
                    newstr = str.partition('\n')
                    str = newstr[0]+' '+newstr[2]
                if j == 0:
                    sqlLine = al.text('select "ID" from "EntityProp" where "Name" = :st ;')
                    try:
                        q = con.execute(sqlLine, {"st": str})
                    except:
                        epID = -1
                        print('(epID) =',epID)
                    else:
                        for raw in q:
                            epID = raw[0]
                            #print(epID)

                if len(str) > 20:
                    str =  str.ljust(70,' ')

                if len(str) < 5:
                    str = str.ljust(7, ' ')

                outfile.write(str + '\t\t')

                if j == 2:
                    if str.find(',') > 0:
                        com2point = str.partition(',')
                        value = com2point[0]+'.'+com2point[2]
                    else:
                        value = str



            j += 1

        sqlLine = al.text('insert into "Chim" values (:tm, :ep, :vl)')
        #print (j,dateStr, epID, value)

        try:
            if int(epID) > 0:
                #q = con.execute(sqlLine, {"tm": dateStr,"ep": int(epID), "vl": float(value)})
                q = con.execute(sqlLine, {"tm": lin,"ep": int(epID), "vl": float(value)})
            else:
                print('Completed.')
        except:
            print('Запись в базу не удалась!')
        i += 1
        outfile.write('\n')


